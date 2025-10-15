import getpass
import re
import subprocess
from pathlib import Path
import os
import platform
import sys

# --- Safe Import for Windows-only modules ---
try:
    if platform.system() == "Windows":
        import win32api
        import win32net
    else:
        win32api = None
        win32net = None
except ImportError:
    win32api = None
    win32net = None

from docxtpl import DocxTemplate
import docx


class commonMethods:
    def get_project_root(self):
        return str(Path(__file__).parent.parent)

    """
    Below code is when Microsoft Word is not installed; 
    we can use python-docx-template to create one.
    """

    def create_docx_file_withoutMS(self):
        doc = DocxTemplate(os.path.join("template.docx"))
        context = {
            "name": "Vincent",
            "company": "WinWire Technologies",
            "location": "Bangalore"
        }

        doc.render(context)
        temp_dir = os.getenv("TEMP", "/tmp")
        filename = os.path.join(temp_dir, "test.docx")
        doc.save(filename)
        print(f"File saved to: {filename}")
        return filename

    def create_docx_file(self):
        filename = ""
        try:
            doc = docx.Document()
            doc.add_paragraph("Hello, World!")
            doc.add_paragraph("This is a test document.")

            temp_dir = os.getenv("TEMP", "/tmp")
            filename = os.path.join(temp_dir, "test.docx")
            doc.save(filename)
            print(f"File saved to: {filename}")
        except Exception as e:
            print(f"Error in create_docx_file -> {e}")
        return filename

    def eTMF_docx_file(self):
        base_filename = "TestAutomation"
        extension = ".docx"
        try:
            doc = docx.Document()
            doc.add_paragraph("Hello, World!")
            doc.add_paragraph("This is a test document.")

            temp_dir = os.getenv("TEMP", "/tmp")
            counter = 1
            while True:
                filename = os.path.join(temp_dir, f"{base_filename}{counter}{extension}")
                if not os.path.exists(filename):
                    break
                counter += 1

            doc.save(filename)
            print(f"File saved to: {filename}")
        except Exception as e:
            print(f"Error in eTMF_docx_file -> {e}")
        return filename

    def error_message(self, message):
        try:
            lines = re.split("[{(|)]", message)
            return lines[0] if lines else message
        except Exception as e:
            print(f"Error in error_message -> {e}")
            return message

    def getusername(self):
        full_name = ""
        try:
            if platform.system() == 'Windows' and win32api and win32net:
                username = win32api.GetUserName()
                user_info = win32net.NetUserGetInfo(None, username, 2)
                full_name = user_info.get('full_name', username)
            else:
                username = getpass.getuser()
                try:
                    import pwd
                    full_name = pwd.getpwnam(username).pw_gecos.split(',')[0]
                except (ImportError, KeyError):
                    full_name = username

            if not full_name:
                full_name = os.getenv("USER", "UnknownUser")
        except Exception as e:
            print("CommonMethods -> getusername -> \n" + str(e))
            if not full_name:
                full_name = os.getenv("USER", "UnknownUser")
        return full_name

    def getOS(self):
        try:
            os_name = sys.platform
            if os_name.startswith("linux"):
                return "linux"
            elif os_name == "darwin":
                return "mac"
            elif os_name == "win32":
                return "windows"
            else:
                raise Exception("OS not supported")
        except Exception as e:
            print("CommonMethods -> getOS -> \n " + str(e))
            return "unknown"

    def get_os_details(self):
        os_name = platform.system()

        if os_name == "Windows":
            windows_version = platform.release()
            version_map = {
                "10": "Windows 10",
                "11": "Windows 11",
                "8": "Windows 8",
                "8.1": "Windows 8.1",
                "7": "Windows 7",
            }
            windows_name = version_map.get(windows_version, "Windows")
            return f"{windows_name} {platform.version()}"

        elif os_name in ["Darwin", "Linux"]:
            try:
                mac_version = subprocess.run(
                    ["sw_vers", "-productVersion"],
                    capture_output=True, text=True
                ).stdout.strip()
            except Exception:
                mac_version = platform.release()

            macos_version_map = {
                "14": "macOS Sonoma",
                "13": "macOS Ventura",
                "12": "macOS Monterey",
                "11": "macOS Big Sur",
                "10.15": "macOS Catalina",
                "10.14": "macOS Mojave",
            }
            major_version = mac_version.split('.')[0]
            mac_name = macos_version_map.get(major_version, os_name)
            return f"{mac_name} {mac_version}"

        else:
            return "Operating system details not available."
